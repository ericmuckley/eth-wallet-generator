import os
import json
import qrcode
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH




with open('wallet_info.json') as f:
    wallet_info = json.load(f)

wallets = wallet_info["wallets"][:57]



doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'

section = doc.sections[0]

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')





for i, wallet in enumerate(wallets):


    print(f'i={i} {wallet["address"]}')
    etherscan_url = f"https://etherscan.io/address/{wallet['address']}#tokentxnsErc721"


    # save etherscan link to address as QR code
    qrcode_etherscan_fp = os.path.join("qr-codes-etherscan", f"{str(i).zfill(3)}.png")
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=2,
    )
    qr.add_data(etherscan_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(qrcode_etherscan_fp)




    current_section = doc.sections[0]  
    current_section.start_type


    p = doc.add_paragraph()
    run = p.add_run("Title Message\n")
    run.font.name = 'Arial'
    run.font.size = Pt(26)
    run.bold = True


    p = doc.add_paragraph()

    run = p.add_run("\n\nYour wallet's secret private key:\n")
    run.bold = True
    run = p.add_run(wallet['private_key'])
    run.font.size = Pt(10)

    run = p.add_run("\n\nYour wallet's secret 24 seed words:\n")
    run.bold = True
    run = p.add_run(wallet['words'])
    run.italic = True



    p = doc.add_paragraph()
    run = p.add_run(f"Your wallet's public address:\n{wallet['address']}\n")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run = p.add_run("This is what other people see when you send and receive cryptocurrencies and NFTs.")
    run.font.name = 'Arial'
    run.font.size = Pt(12)


    p = doc.add_paragraph()
    run = p.add_run(" ")
    run = p.add_run("\n")
    

    p = doc.add_paragraph()
    run = p.add_run("View the full history of your wallet using the QR code below. ")
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    img = doc.add_picture(qrcode_etherscan_fp, width=Inches(2.5))
    p = doc.add_paragraph()
    run = p.add_run(" ")


    doc.add_page_break()

doc.save("exported_wallet_details.docx")

